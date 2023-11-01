# this version compared the dates of the letters to create a new section for each date.
# I'm no longer doing this, but I'm keeping it here for reference.

img_path = "C:\\Git\\read-letters\\letters"

# initialize
previous_date = None

import docx
document = docx.Document()

import os
# find the file names from the directory.  
for path,dirs,files in os.walk(img_path):
    for file in files: 
        if file.split('.')[1] == 'jpg': #only process the image files.  there are corresponding txt files for every jpg file.
            # Remove the extension and split the filename by '-'
            parts = file.split('.')[0].split('-')
            # Assign the first three parts to 'date' and the fourth part to 'where'
            date = '-'.join(parts[:3])
            who = parts[3].capitalize()  
            p = parts[4]
            where = parts[5].capitalize() 
            heading = f"From {who}, Postmarked: {date} {where}"
            page = f"p {p}"
            # find the text file for this image
            txt_file = f"{file.split('.')[0]}.txt"
            txt_path = os.path.join(path,txt_file)
            img_path = os.path.join(path,file)
            with open(txt_path) as f:
                txt = f.read()
            f.close()

            if previous_date is None:
                # first time through 
                document.add_heading(heading, level=2)
                # Add a header to the new section
                # add first table
                table = document.add_table(rows=1, cols=2)
                row_cells = table.add_row().cells

                # Add a picture to the first cell
                run = row_cells[0].paragraphs[0].add_run()
                run.add_picture(img_path,width=docx.shared.Inches(4))

                # Add text to the second cell
                row_cells[1].text = txt


            else:
                # Compare with the previous file
                if date == previous_date:
                    row_cells = table.add_row().cells
                    a, b = row_cells[:2]
                    a.merge(b)
                    row_cells[0].text = f"{heading}, {page}"
                    row_cells[0].paragraphs[0].paragraph_format.keep_with_next = True
                    # add the image and text for this file to the document
                    row_cells = table.add_row().cells
                    # Add a picture to the first cell
                    run = row_cells[0].paragraphs[0].add_run()
                    run.add_picture(img_path,width=docx.shared.Inches(4))
                    # Add text to the second cell
                    row_cells[1].text = txt

                else:
                    # Start a new section for this date
                    document.add_heading(heading, level=2) 
                    # add a new table
                    table = document.add_table(rows=1, cols=2)
                    row_cells = table.add_row().cells
                    # Add a picture to the first cell
                    run = row_cells[0].paragraphs[0].add_run()
                    run.add_picture(img_path,width=docx.shared.Inches(4))
                    # Add text to the second cell
                    row_cells[1].text = txt



            # each time, update the previous date
            previous_date = date

document.save("letters.docx")   