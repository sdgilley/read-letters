'''  
This script creates a word document & pdf containing letter image and text, 
in a table with two columns.  The first column contains the image and
the second column contains the text.

Before you create the doc, use read-letters.py to create the text files for each image.
Manually edit the text files if necessary.  
Text and image files should be in the same folder.
This script assumes the following structure for the image names:

mm-dd-yy-who-p-where.jpg (or png)
i.e.:    01-11-54-dad-1-richmond.jpg

Install these packages before running this script:
      pip install python-docx
      pip install docx2pdf
'''

##### USER INPUT #####
# path to the folder with image a text files.
img_path = "C:\\Git\\read-letters-data\\letters" 
title = "Letters from Mom & Dad"
century = "19" # century of the letter
docname = "letters.docx" #name of the output document
######################

import docx
from docx.shared import Inches
from docx2pdf import convert

# create a document and set up margins
document = docx.Document()
sections = document.sections
for section in sections:
    section.left_margin = Inches(.75)
    section.right_margin = Inches(.75)
# add a title
document.add_heading(title, 0)
# add a table
table = document.add_table(rows=1, cols=2)
pnum = 0

# iterate through the files in the folder.  
import os
# find the file names from the directory.  
for path,dirs,files in os.walk(img_path):
    for file in sorted(files): 
        if file.split('.')[1] in ['jpg', 'png']: # only process the image files.
            # Remove the extension and split the filename by '-'
            parts = file.split('.')[0].split('-')

            # split apart name to get the date, who, where, and page number.
            date = '/'.join(parts[:2])
            year = f"{century}{parts[2]}" # add century to year
            date = f"{date}/{year}"
            who = parts[3].capitalize()  
            p = parts[4]
            where = parts[5].capitalize() 
            # create a heading for this particular letter
            heading = f"From {who}, Postmarked {date} {where}, p {p}"

            # construct full file paths for both text and image files
            txt_file = f"{file.split('.')[0]}.txt"
            txt_path = os.path.join(path,txt_file)
            img_path = os.path.join(path,file)

            # read the text file for this image
            with open(txt_path) as f:
                txt = f.read()
            f.close()
            
            # add a merged cell with info about this page.  
            row_cells = table.add_row().cells
            a, b = row_cells[:2]
            a.merge(b)
            row_cells[0].text = f"{heading}"
            # keep this cell with the image/text row below
            row_cells[0].paragraphs[0].paragraph_format.keep_with_next = True

            # add the image and text 
            row_cells = table.add_row().cells
            # Add image to the first cell
            run = row_cells[0].paragraphs[0].add_run()
            run.add_picture(img_path,width=Inches(4))
            # Add text to the second cell
            row_cells[1].text = txt
            pnum += 1

document.save(docname)   
pdfname = docname.split('.')[0] + ".pdf"
convert(docname, pdfname)

print("\n**********************************************************")
print(f"Created {docname} & {pdfname} with {pnum} letters pages.")